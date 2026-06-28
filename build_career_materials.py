from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("career_materials")
OUT.mkdir(exist_ok=True)


PROFILE = {
    "name": "赵玺",
    "phone": "15306093416",
    "email": "zhaoxi09@hotmail.com",
    "location": "上海市普陀区",
    "availability": "可随时到岗；可连续实习6个月及以上",
}


EDUCATION = [
    {
        "time": "2025.09-2028.07",
        "school": "华东师范大学",
        "degree": "硕士",
        "major": "应用心理（人因工程与用户体验方向）",
        "details": [
            "CGPA: 3.69/4.0；研究方向：智能交通系统/智能汽车中的人因工程",
            "核心课程：用户体验与可用性研究、人因工程设计与评价、深度学习",
        ],
    },
    {
        "time": "2019.09-2023.07",
        "school": "厦门大学",
        "degree": "本科",
        "major": "计算机科学与技术",
        "details": [
            "CGPA: 3.2/4.0；核心课程：计算机架构、算法设计与分析、多媒体技术",
        ],
    },
]


WORK = [
    {
        "time": "2024.03-2024.12",
        "org": "厦门五五安科信息科技有限公司",
        "role": "产品工程师",
        "bullets": [
            "参与PC端电子取证软件系统重构，负责需求梳理、竞品与业务流程分析、原型设计及研发协同推进。",
            "将业务规则、操作流程与用户痛点转化为可执行的产品方案，支持核心软件从需求整理到上线落地。",
            "在市场调研、需求分析、产品规划、跨部门沟通和用户体验优化中积累完整产品交付经验。",
        ],
    }
]


PROJECTS = {
    "phenodrive_hf": {
        "title": "PhenoDrive：面向里程焦虑的智能电动汽车自适应界面与驾驶员压力画像研究",
        "time": "2026.01-至今",
        "role": "独立主导",
        "bullets": [
            "围绕智能电动汽车里程焦虑场景，设计模拟驾驶高压力诱发实验，采集眼动、行为序列等多维数据，探索客观表型驱动的自适应HMI干预路径。",
            "完成压力画像挖掘与自适应HMI设计核心工作，并以UIST'26投稿为目标整理研究叙事、实验流程和交互方案。",
            "开展多模态生理计算与动态信任重建验证，计划结合EEG、眼动与隐马尔可夫模型分析用户隐状态转移和关键干预时间窗。",
        ],
    },
    "speech_ai": {
        "title": "车载语音情感对话链路校企合作项目",
        "time": "2025.12-2026.06（日期待确认）",
        "role": "研究与评测支持",
        "bullets": [
            "基于真实车载对话场景整理情感对话Benchmark，构建覆盖情绪识别、上下文理解与回应质量的评测材料。",
            "采用LLM-as-a-judge范式进行效果验证，辅助形成可复用的评测流程和结果分析口径。",
            "建议对外表述仅保留问题域、方法和非敏感产出，不公开合作企业、原始语料、具体业务场景和未授权指标。",
        ],
    },
    "social_anxiety": {
        "title": "社交焦虑训练平台（个人vibe coding项目）",
        "time": "2026",
        "role": "产品设计与原型开发",
        "bullets": [
            "面向社交焦虑情境训练，设计从情境选择、对话练习、反馈复盘到进阶训练的完整产品流程。",
            "使用AI辅助编程快速完成可交互原型，验证面向心理训练场景的任务流、反馈机制和内容生成策略。",
            "体现从用户问题定义、交互结构设计到前端原型落地的端到端产品验证能力。",
        ],
    },
    "course_research": {
        "title": "游戏领域用户研究报告与UI/交互设计课程项目",
        "time": "课程项目",
        "role": "用户研究/交互设计",
        "bullets": [
            "完成游戏用户研究报告，覆盖研究问题定义、用户洞察、体验问题归纳和设计建议输出。",
            "完成一套UI与交互设计方案，包含信息架构、关键页面、交互流程和视觉规范，可作为网站作品集的过程型案例。",
        ],
    },
    "leapmotion": {
        "title": "本科毕设：基于Unity与Leap Motion的交互系统开发",
        "time": "本科毕设",
        "role": "独立开发",
        "bullets": [
            "使用Unity与Leap Motion完成手势交互系统开发，覆盖交互逻辑、场景搭建和可用性验证。",
            "适合作为早期HCI/交互技术能力证明，网站中建议突出问题场景、交互机制和演示视频/截图。",
        ],
    },
}


SKILLS_COMMON = [
    "研究方法：用户访谈、可用性测试、实验设计、问卷与量表、竞品/文献/案头研究、报告撰写",
    "数据与工具：眼动、EEG、fNIRS、心率等人因数据采集流程；Excel/SPSS/Python基础分析能力",
    "设计与产品：Figma、Adobe Illustrator、Photoshop；需求分析、原型设计、产品文档、跨部门协同",
    "AI与开发：理解深度学习、大模型推理与AI Agent基础概念；能够使用AI工具快速搭建产品原型",
    "英语：IELTS 7.0（听力7.5，阅读8.0，写作6.5，口语6.5）",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, bold=False, size=9.5, color="222222"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=10.5, color="1F4E79")
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    p_bdr.append(bottom)
    p._p.get_or_add_pPr().append(p_bdr)


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("• ")
    set_font(r, size=9.2, color="1F4E79")
    r = p.add_run(text)
    set_font(r, size=9.2)


def add_entry(doc, time, title, subtitle=None, bullets=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(time + "  ")
    set_font(r, bold=True, size=9.3)
    r = p.add_run(title)
    set_font(r, bold=True, size=9.5)
    if subtitle:
        r = p.add_run("  |  " + subtitle)
        set_font(r, size=9.2, color="555555")
    for b in bullets or []:
        add_bullet(doc, b)


def build_resume(kind):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(PROFILE["name"])
    set_font(r, bold=True, size=18, color="1F4E79")
    title.paragraph_format.space_after = Pt(1)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f'{PROFILE["phone"]} | {PROFILE["email"]} | {PROFILE["location"]} | {PROFILE["availability"]}')
    set_font(r, size=8.8, color="444444")
    info.paragraph_format.space_after = Pt(4)

    intent = "人因工程/用户研究实习生" if kind == "hf" else "AI产品/交互产品实习生"
    summary = (
        "应用心理人因工程与用户体验方向硕士在读，计算机科学本科背景；具备实验设计、体验量化、眼动/EEG等人因数据采集与用户研究报告能力，"
        "并有产品工程、交互设计和智能汽车HMI研究经历。"
        if kind == "hf"
        else
        "应用心理人因工程与用户体验方向硕士在读，计算机科学本科背景；具备AI产品原型验证、需求分析、用户研究、交互设计和LLM评测经验，"
        "能在研究洞察、产品方案和可交互原型之间快速闭环。"
    )
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    widths = [Cm(2.5), Cm(15.8)]
    labels = ["求职方向", "个人定位"]
    values = [intent, summary]
    for i in range(2):
        cells = table.rows[i].cells
        cells[0].width = widths[0]
        cells[1].width = widths[1]
        set_cell_shading(cells[0], "EAF2F8")
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(labels[i])
        set_font(r, bold=True, size=9.2, color="1F4E79")
        cells[1].text = ""
        r = cells[1].paragraphs[0].add_run(values[i])
        set_font(r, size=9.2)

    add_heading(doc, "教育背景")
    for edu in EDUCATION:
        add_entry(doc, edu["time"], f'{edu["school"]}  {edu["degree"]}  {edu["major"]}', bullets=edu["details"])

    add_heading(doc, "工作经历")
    for item in WORK:
        add_entry(doc, item["time"], f'{item["org"]}  {item["role"]}', bullets=item["bullets"])

    add_heading(doc, "项目经历")
    if kind == "hf":
        order = ["phenodrive_hf", "speech_ai", "course_research", "leapmotion"]
    else:
        order = ["social_anxiety", "speech_ai", "phenodrive_hf", "course_research"]
    for key in order:
        item = PROJECTS[key]
        add_entry(doc, item["time"], item["title"], item["role"], item["bullets"])

    add_heading(doc, "技能与荣誉")
    add_bullet(doc, "荣誉：2025年全国应用心理实践技能大赛产品设计赛道三等奖")
    for skill in SKILLS_COMMON:
        add_bullet(doc, skill)

    for section in doc.sections:
        section.start_type = WD_SECTION.CONTINUOUS

    path = OUT / ("赵玺_人因用户研究方向_简历初稿.docx" if kind == "hf" else "赵玺_AI产品方向_简历初稿.docx")
    doc.save(path)
    return path


def write_strategy():
    text = """# 个人网站内容策略

## 简历版本建议

建议制作两版简历，而不是一份通投简历。

- 人因工程/用户研究版：主线是实验设计、体验量化、用户研究、HMI、眼动/EEG等客观数据能力。优先投人因研究工程师、用户研究员、体验研究、智能座舱体验等岗位。
- AI产品版：主线是AI产品理解、Agent/LLM评测、产品原型、交互设计、产品工程经历。优先投AI产品实习、智能硬件产品、具身智能多模态交互、AI Agent相关岗位。

两版简历共享教育、工作经历和核心项目，但项目排序、标题关键词和个人定位不同。

## 网站栏目

1. 首页：一句定位 + 两个入口：Human Factors & UX Research / AI Product & Prototyping。
2. Featured Research：PhenoDrive。展示研究问题、方法、交互框架、可公开结果；UIST'26视频先不公开视频，使用密码保护、私密链接或“available upon request”。
3. Industry Collaboration：车载语音情感对话链路。展示问题域、方法框架、评测范式和脱敏产出，不展示企业名称、原始语料、未授权Benchmark细节和内部指标。
4. AI Product / Vibe Coding：社交焦虑训练平台。展示用户痛点、训练流程、核心交互、技术实现和反思。
5. UX & Design Coursework：游戏领域用户研究报告、完整UI+交互设计。适合放过程图、研究洞察、设计迭代和最终界面。
6. HCI Prototype Archive：本科毕设Unity+Leap Motion项目，以及本科网页开发、UI设计等小项目。作为能力补充，不放在首页最重位置。
7. About / Contact：教育背景、方法栈、工具栈、邮箱。

## 公开与脱敏规则

- 投稿中的UIST'26视频：在录用或明确允许公开前，不建议放公开视频。可放项目摘要、静态图、模糊化流程图和“demo available upon request”。
- 校企合作项目：默认脱敏。公司名称、真实用户数据、原始对话、内部Benchmark、未公开指标、业务策略都不放；只展示你的方法、流程、角色和泛化产出。
- 简历：可以写合作项目和方法，但避免写“企业那边应该没关系”这类不确定表述。用“校企合作项目”“真实车载场景”“脱敏Benchmark/评测流程”即可。
- 网站：公开站点只放你有权公开的截图、流程和结论；敏感材料可以放在私密PDF或面试现场展示版。
"""
    path = OUT / "个人网站与简历策略.md"
    path.write_text(text, encoding="utf-8")
    return path


if __name__ == "__main__":
    paths = [build_resume("hf"), build_resume("ai"), write_strategy()]
    for p in paths:
        print(p.resolve())
